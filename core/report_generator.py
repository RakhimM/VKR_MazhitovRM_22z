from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class ReportGenerator:
    @staticmethod
    def generate(filepath, data):
        doc = Document()
        results = data["results"]

        title = doc.add_heading("Отчёт о тестировании", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Дата и время: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        if data["include_ram"]:
            doc.add_heading("1. Тест внешнего ОЗУ", level=1)
            ram = results.get("ram", {})
            doc.add_paragraph(f"Количество ошибок: {ram.get('errors', 0)}")
            if ram.get("details"):
                doc.add_paragraph("Подробности:")
                for detail in ram["details"]:
                    doc.add_paragraph(f"  • {detail}", style='List Bullet')
            doc.add_paragraph()

        if data["include_rom"]:
            doc.add_heading("2. Тест ПЗУ (CRC32)", level=1)
            rom = results.get("rom", {})
            doc.add_paragraph(f"Ошибок: {rom.get('errors', 0)}")
            doc.add_paragraph(f"Вычисленная CRC32: 0x{rom.get('crc', 0):08X}")
            expected = rom.get("expected_crc")
            if expected is not None:
                doc.add_paragraph(f"Ожидаемая CRC32: 0x{expected:08X}")
                if rom.get('crc') == expected:
                    doc.add_paragraph("Статус: ПЗУ в порядке", style='Intense Quote')
                else:
                    doc.add_paragraph("Статус: НЕСООТВЕТСТВИЕ CRC!", style='Intense Quote')
            doc.add_paragraph()

        if data["include_exc"]:
            doc.add_heading("3. Тесты исключений CPU/FPU", level=1)
            exc = results.get("exc", {})
            doc.add_paragraph(f"Количество ошибок: {exc.get('errors', 0)}")
            if exc.get("details"):
                doc.add_paragraph("Детали:")
                for detail in exc["details"]:
                    doc.add_paragraph(f"  • {detail}", style='List Bullet')
            doc.add_paragraph()

        if data["include_summary"]:
            doc.add_heading("4. Общая сводка", level=1)
            summary = results.get("summary", {})
            doc.add_paragraph(f"Выполнено циклов: {summary.get('cycles', 0)}")
            doc.add_paragraph(f"Всего ошибок: {summary.get('total_errors', 0)}")
            doc.add_paragraph()

        if data.get("comment"):
            doc.add_heading("5. Комментарий", level=1)
            doc.add_paragraph(data["comment"])

        doc.save(filepath)